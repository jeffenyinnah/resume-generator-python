import streamlit as st
from openai import OpenAI
from docx import Document
import os
from dotenv import load_dotenv
from pymongo import MongoClient
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

# Load environment variables
load_dotenv()

# Set OpenAI API key
client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

# MongoDB Atlas connection
MONGO_URI = os.getenv("MONGO_URI")
DB_NAME = "resume_generator"
COLLECTION_NAME = "users"

# Connect to MongoDB
def connect_to_mongodb():
    try:
        client = MongoClient(MONGO_URI)
        db = client[DB_NAME]
        collection = db[COLLECTION_NAME]
        return collection
    except Exception as e:
        st.error(f"Erro ao conectar ao MongoDB: {e}")
        return None

# Function to check if a user exists in the database
def check_user_exists(email, collection):
    user = collection.find_one({"email": email})
    return user is not None

# Function to authenticate a user
def authenticate_user(email, password, collection):
    user = collection.find_one({"email": email, "password": password})
    if user:
        return user
    else:
        return None

# Function to add a new user to the database
def add_user(email, password, collection):
    try:
        user_data = {"email": email, "password": password, "generation_count": 0}
        collection.insert_one(user_data)
        st.success("Cadastro bem-sucedido!")
    except Exception as e:
        st.error(f"Erro ao cadastrar usuário: {e}")

# Function to apply the template's styles
def apply_template_styles(doc, template_doc):
    """Copy styles from template to the new document"""
    # Copy the styles from the template
    for style in template_doc.styles:
        if style.name not in doc.styles:
            doc.styles.add_style(style.name, style.type)
    
    # Try to match paragraph formatting
    for i, paragraph in enumerate(doc.paragraphs):
        try:
            # If there's a matching paragraph in the template, copy its style
            if i < len(template_doc.paragraphs):
                template_paragraph = template_doc.paragraphs[i]
                paragraph.style = template_paragraph.style
                
                # Copy paragraph formatting
                paragraph_format = paragraph.paragraph_format
                template_format = template_paragraph.paragraph_format
                
                # Copy attributes
                paragraph_format.alignment = template_format.alignment
                if template_format.line_spacing:
                    paragraph_format.line_spacing = template_format.line_spacing
                if template_format.space_before:
                    paragraph_format.space_before = template_format.space_before
                if template_format.space_after:
                    paragraph_format.space_after = template_format.space_after
        except Exception as e:
            print(f"Error copying style for paragraph {i}: {e}")
    
    return doc

# Function to generate resume content using OpenAI
def generate_resume(name, email, phone, industry, job_type, experiences, educations, skills, languages, linkedin):
    prompt = f"""
    Crie um currículo profissional em português para {name}, que está buscando uma vaga de {job_type} na indústria de {industry}.
    
    **Instruções Específicas:**
    1. **Perfil:**
       - Escreva um breve perfil profissional (2-3 frases) que destaque as habilidades e experiências mais relevantes para o cargo de {job_type}.
       - NÃO use marcações de markdown como ** para negrito.

    2. **Experiências Profissionais:**
       - Formate as experiências no formato "cargo | empresa | período (mês ano - mês ano)"
       - Cada responsabilidade deve começar com um verbo no passado e ser escrita em tópicos (bullet points com "-")
       - Inclua 3-4 responsabilidades por experiência
       - Destaque conquistas quantificáveis (ex: "Aumentou as vendas em 20%", "Reduziu custos em 15%")
       - NÃO use marcações de markdown como ** para negrito.

    3. **Educação:**
       - Formate como "grau | mês ano | instituição, cidade, estado"
       - NÃO use marcações de markdown como ** para negrito.

    4. **Habilidades:**
       - Liste as habilidades como tópicos separados
       - Mantenha cada item breve e direto
       - NÃO use marcações de markdown como ** para negrito.

    5. **Idiomas:**
       - Liste todos os idiomas fornecidos
       - NÃO use marcações de markdown como ** para negrito.

    6. **Atividades e Interesses:**
       - Liste interesses separados por vírgulas em uma única linha
       - NÃO use marcações de markdown como ** para negrito.

    IMPORTANTE: NÃO USE NENHUMA FORMATAÇÃO MARKDOWN COMO **, ##, OU OUTROS SÍMBOLOS DE FORMATAÇÃO. O TEXTO DEVE SER SIMPLES SEM MARCAÇÕES.

    **Detalhes do Candidato:**
    - Nome: {name}
    - Email: {email}
    - Telefone: {phone}
    - LinkedIn: {linkedin}
    - Idiomas: {', '.join(languages)}

    **Experiências Profissionais:**
    {experiences}

    **Educação:**
    {educations}

    **Habilidades:**
    {', '.join(skills)}
    """
    
    response = client.chat.completions.create(
        model="gpt-4",
        messages=[
            {"role": "system", "content": "Você é um especialista em redação de currículos profissionais. Crie currículos com conteúdo conciso e impactante usando texto simples sem formatação markdown."},
            {"role": "user", "content": prompt}
        ]
    )
    return response.choices[0].message.content

# Function to list available templates
def list_templates():
    templates_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "templates")  # Use absolute path
    if not os.path.exists(templates_dir):
        os.makedirs(templates_dir)
        st.warning(f"Diretório 'templates' criado em: {os.path.abspath(templates_dir)}")
    templates = [f for f in os.listdir(templates_dir) if f.endswith(".docx")]
    if not templates:
        st.warning("Nenhum template encontrado na pasta 'templates'. Adicione templates .docx para continuar.")
    return templates, templates_dir

# Function to load the selected template
def load_template(template_name, templates_dir):
    template_path = os.path.join(templates_dir, template_name)
    print(f"Tentando carregar o template: {template_path}")  # Debugging
    if not os.path.exists(template_path):
        st.error(f"Arquivo não encontrado: {template_path}")
        return None
    try:
        doc = Document(template_path)
        return doc
    except Exception as e:
        st.error(f"Erro ao carregar o template: {e}")
        return None

# Function to format template2
def template2(doc, resume_content, name, email, phone, linkedin):
    # Clear the document
    for paragraph in doc.paragraphs[:]:
        p = paragraph._element
        p.getparent().remove(p)
        paragraph._p = paragraph._element = None
    
    # Add contact info at top
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_run = header.add_run(name)
    header_run.bold = True
    header_run.font.size = Pt(16)
    
    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.add_run(f"{phone} | {email}")
    if linkedin:
        contact.add_run(f" | {linkedin}")
    
    # Add sections with proper ATS formatting
    sections = [
        ("PERFIL", "Perfil"),
        ("EXPERIÊNCIA", "Experiência Profissional"),
        ("EDUCAÇÃO", "Educação"),
        ("HABILIDADES E COMPETÊNCIAS", "Habilidades"),
        ("IDIOMAS", "Idiomas"),
        ("ATIVIDADES E INTERESSES", "Atividades e Interesses")
    ]
    
    for section_title, content_key in sections:
        # Add section header
        section_header = doc.add_paragraph()
        section_header.style = 'Heading 1'
        section_run = section_header.add_run(section_title)
        section_run.bold = True
        section_run.font.all_caps = True
        
        # Get content for this section
        content = ""
        if content_key == "Perfil":
            if "Experiência Profissional" in resume_content:
                content = resume_content.split("Experiência Profissional")[0].strip()
            else:
                content = resume_content
        elif content_key == "Experiência Profissional":
            if content_key in resume_content and "Educação" in resume_content:
                content = resume_content.split(content_key)[1].split("Educação")[0].strip()
        elif content_key == "Educação":
            if content_key in resume_content and "Habilidades" in resume_content:
                content = resume_content.split(content_key)[1].split("Habilidades")[0].strip()
        elif content_key == "Habilidades":
            if content_key in resume_content:
                if "Idiomas" in resume_content:
                    content = resume_content.split(content_key)[1].split("Idiomas")[0].strip()
                else:
                    content = resume_content.split(content_key)[1].strip()
        elif content_key == "Idiomas":
            if content_key in resume_content:
                if "Atividades e Interesses" in resume_content:
                    content = resume_content.split(content_key)[1].split("Atividades e Interesses")[0].strip()
                else:
                    content = resume_content.split(content_key)[1].strip()
        elif content_key == "Atividades e Interesses":
            if content_key in resume_content:
                content = resume_content.split(content_key)[1].strip()
        
        # Format the content appropriately
        if content:
            if content_key == "Habilidades":
                skills_list = content.split(",")
                for skill in skills_list:
                    skill = skill.strip()
                    if skill:
                        doc.add_paragraph(skill, style='List Bullet')
            elif content_key == "Experiência Profissional":
                # Parse experience entries
                experience_entries = content.split("\n\n")
                for entry in experience_entries:
                    lines = entry.strip().split("\n")
                    if lines and len(lines) >= 1:
                        title_line = lines[0].strip()
                        if " | " in title_line:
                            parts = title_line.split(" | ")
                            if len(parts) >= 3:
                                job_title, company, dates = parts[0], parts[1], parts[2]
                                job_heading = doc.add_paragraph()
                                job_heading.add_run(f"{job_title.upper()} | {company.upper()} | {dates.upper()}").bold = True
                        
                        # Add responsibilities as bullet points
                        for i in range(1, len(lines)):
                            line = lines[i].strip()
                            if line.startswith("-"):
                                p = doc.add_paragraph(line[1:].strip(), style='List Bullet')
                            else:
                                doc.add_paragraph(line)
            else:
                # Add regular paragraphs for other sections
                paragraphs = content.split("\n")
                for para in paragraphs:
                    if para.strip():
                        doc.add_paragraph(para.strip())
    
    return doc

# Function to format template 1
def template1(doc, resume_content, name, email, phone, linkedin):
    # Clear the document
    for paragraph in doc.paragraphs[:]:
        p = paragraph._element
        p.getparent().remove(p)
        paragraph._p = paragraph._element = None
    
    # Add contact info at top
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_run = header.add_run(name)
    header_run.bold = True
    header_run.font.size = Pt(16)
    
    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.add_run(f"{phone} | {email}")
    if linkedin:
        contact.add_run(f" | {linkedin}")
    
    # Add sections with proper ATS formatting
    sections = [
        ("PERFIL", "Perfil"),
        ("EXPERIÊNCIA", "Experiência Profissional"),
        ("EDUCAÇÃO", "Educação"),
        ("HABILIDADES E COMPETÊNCIAS", "Habilidades"),
        ("IDIOMAS", "Idiomas"),
        ("ATIVIDADES E INTERESSES", "Atividades e Interesses")
    ]
    
    for section_title, content_key in sections:
        # Add section header
        section_header = doc.add_paragraph()
        section_header.style = 'Heading 1'
        section_run = section_header.add_run(section_title)
        section_run.bold = True
        section_run.font.all_caps = True
        
        # Get content for this section
        content = ""
        if content_key == "Perfil":
            if "Experiência Profissional" in resume_content:
                content = resume_content.split("Experiência Profissional")[0].strip()
            else:
                content = resume_content
        elif content_key == "Experiência Profissional":
            if content_key in resume_content and "Educação" in resume_content:
                content = resume_content.split(content_key)[1].split("Educação")[0].strip()
        elif content_key == "Educação":
            if content_key in resume_content and "Habilidades" in resume_content:
                content = resume_content.split(content_key)[1].split("Habilidades")[0].strip()
        elif content_key == "Habilidades":
            if content_key in resume_content:
                if "Idiomas" in resume_content:
                    content = resume_content.split(content_key)[1].split("Idiomas")[0].strip()
                else:
                    content = resume_content.split(content_key)[1].strip()
        elif content_key == "Idiomas":
            if content_key in resume_content:
                if "Atividades e Interesses" in resume_content:
                    content = resume_content.split(content_key)[1].split("Atividades e Interesses")[0].strip()
                else:
                    content = resume_content.split(content_key)[1].strip()
        elif content_key == "Atividades e Interesses":
            if content_key in resume_content:
                content = resume_content.split(content_key)[1].strip()
        
        # Format the content appropriately
        if content:
            if content_key == "Habilidades":
                skills_list = content.split(",")
                for skill in skills_list:
                    skill = skill.strip()
                    if skill:
                        doc.add_paragraph(skill, style='List Bullet')
            elif content_key == "Experiência Profissional":
                # Parse experience entries
                experience_entries = content.split("\n\n")
                for entry in experience_entries:
                    lines = entry.strip().split("\n")
                    if lines and len(lines) >= 1:
                        title_line = lines[0].strip()
                        if " | " in title_line:
                            parts = title_line.split(" | ")
                            if len(parts) >= 3:
                                job_title, company, dates = parts[0], parts[1], parts[2]
                                job_heading = doc.add_paragraph()
                                job_heading.add_run(f"{job_title.upper()} | {company.upper()} | {dates.upper()}").bold = True
                        
                        # Add responsibilities as bullet points
                        for i in range(1, len(lines)):
                            line = lines[i].strip()
                            if line.startswith("-"):
                                p = doc.add_paragraph(line[1:].strip(), style='List Bullet')
                            else:
                                doc.add_paragraph(line)
            else:
                # Add regular paragraphs for other sections
                paragraphs = content.split("\n")
                for para in paragraphs:
                    if para.strip():
                        doc.add_paragraph(para.strip())
    
    return doc

# Function to format as generic resume
def format_as_generic_resume(doc, resume_content, name, email, phone, linkedin):
     # Clear the document content if it's not empty
    for paragraph in doc.paragraphs[:]:
        p = paragraph._element
        p.getparent().remove(p)
        paragraph._p = paragraph._element = None

    # Add header with contact information
    header_paragraph = doc.add_paragraph()
    header_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header_paragraph.add_run(f"{name}")
    run.bold = True
    run.font.size = Pt(16)
    
    contact_paragraph = doc.add_paragraph()
    contact_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_paragraph.add_run(f"{phone} | {email}")
    if linkedin:
        contact_paragraph.add_run(f" | {linkedin}")

    # Function to clean markdown syntax
    def clean_markdown(text):
        # Remove bold markdown
        text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
        # Remove other markdown formatting as needed
        return text

    # Add Perfil/Profile section
    section_heading = doc.add_paragraph()
    section_heading.add_run("PERFIL").bold = True
    section_heading.style = 'Heading 1'
    
    if "Experiência Profissional" in resume_content:
        profile_text = resume_content.split("Experiência Profissional")[0].strip()
        profile_text = clean_markdown(profile_text)
        doc.add_paragraph(profile_text)
    else:
        profile_text = clean_markdown(resume_content)
        doc.add_paragraph(profile_text)
    
    # Add Experience section
    section_heading = doc.add_paragraph()
    section_heading.add_run("EXPERIÊNCIA").bold = True
    section_heading.style = 'Heading 1'
    
    if "Experiência Profissional" in resume_content and "Educação" in resume_content:
        experience_text = resume_content.split("Experiência Profissional")[1].split("Educação")[0].strip()
        
        # Parse and format experience entries
        experience_entries = experience_text.split("\n\n")
        for entry in experience_entries:
            lines = entry.strip().split("\n")
            if lines and len(lines) >= 1:
                # First line typically contains job title, company, dates
                title_line = clean_markdown(lines[0].strip())
                if " | " in title_line:
                    parts = title_line.split(" | ")
                    if len(parts) >= 3:
                        job_title, company, dates = parts[0], parts[1], parts[2]
                        job_heading = doc.add_paragraph()
                        job_heading.add_run(f"{job_title.upper()} | {company.upper()} | {dates.upper()}").bold = True
                        job_heading.style = 'Heading 2'
                
                # Add bullet points for responsibilities
                for i in range(1, len(lines)):
                    line = clean_markdown(lines[i].strip())
                    if line.startswith("-"):
                        p = doc.add_paragraph(line[1:].strip(), style='List Bullet')
                    else:
                        doc.add_paragraph(line)
    
    # Add Education section
    section_heading = doc.add_paragraph()
    section_heading.add_run("EDUCAÇÃO").bold = True
    section_heading.style = 'Heading 1'
    
    if "Educação" in resume_content and "Habilidades" in resume_content:
        education_text = resume_content.split("Educação")[1].split("Habilidades")[0].strip()
        
        # Parse and format education entries
        education_entries = education_text.split("\n\n")
        for entry in education_entries:
            lines = entry.strip().split("\n")
            if lines and len(lines) >= 1:
                education_paragraph = doc.add_paragraph()
                clean_line = clean_markdown(lines[0])
                education_paragraph.add_run(clean_line.upper()).bold = True
                
                for i in range(1, len(lines)):
                    doc.add_paragraph(clean_markdown(lines[i]))
    
    # Add Skills section
    section_heading = doc.add_paragraph()
    section_heading.add_run("HABILIDADES E COMPETÊNCIAS").bold = True
    section_heading.style = 'Heading 1'
    
    if "Habilidades" in resume_content:
        skills_text = None
        if "Idiomas" in resume_content:
            skills_text = resume_content.split("Habilidades")[1].split("Idiomas")[0].strip()
        else:
            skills_text = resume_content.split("Habilidades")[1].strip()
            
        # Format skills as bullet points
        skills_list = skills_text.split(",")
        for skill in skills_list:
            skill = clean_markdown(skill.strip())
            if skill:
                doc.add_paragraph(skill, style='List Bullet')
    
    # Add Languages section if present
    if "Idiomas" in resume_content:
        section_heading = doc.add_paragraph()
        section_heading.add_run("IDIOMAS").bold = True
        section_heading.style = 'Heading 1'
        
        languages_text = None
        if "Atividades e Interesses" in resume_content:
            languages_text = resume_content.split("Idiomas")[1].split("Atividades e Interesses")[0].strip()
        else:
            languages_text = resume_content.split("Idiomas")[1].strip()
        
        languages_list = languages_text.split(",")
        for language in languages_list:
            language = clean_markdown(language.strip())
            if language:
                doc.add_paragraph(language)
    
    # Add Interests section
    section_heading = doc.add_paragraph()
    section_heading.add_run("ATIVIDADES E INTERESSES").bold = True
    section_heading.style = 'Heading 1'
    
    if "Atividades e Interesses" in resume_content:
        interests_text = resume_content.split("Atividades e Interesses")[1].strip()
        interests_text = clean_markdown(interests_text)
        interests_paragraph = doc.add_paragraph(interests_text)
    
    # Apply consistent formatting
    for paragraph in doc.paragraphs:
        if paragraph.style.name.startswith('Heading'):
            for run in paragraph.runs:
                run.bold = True
                run.font.all_caps = True

# Function to populate the Word document template
def create_word_doc(resume_content, name, email, phone, linkedin, template_name, templates_dir, filename="curriculo.docx"):
     # Load the template for reference
    template_doc = load_template(template_name, templates_dir)
    if template_doc is None:
        st.warning("Template não encontrado. Criando um novo documento do zero.")
        doc = Document()
    else:
        # Create a new document
        doc = Document()
        
        # Based on template name, apply appropriate formatting
        if "template2" in template_name:
            doc = template2(doc, resume_content, name, email, phone, linkedin)
        elif "template1" in template_name:
            doc = template1(doc, resume_content, name, email, phone, linkedin)
        else:
            # Generic formatting
            doc = format_as_generic_resume(doc, resume_content, name, email, phone, linkedin)
        
        # Apply any remaining template styles
        if template_doc:
            doc = apply_template_styles(doc, template_doc)
    
    # Save the document
    doc.save(filename)
    with open(filename, "rb") as file:
        doc_bytes = file.read()
    return doc_bytes

# Function to update generation count for a user
def update_generation_count(email, collection):
    try:
        # Update the generation count in the database
        result = collection.update_one(
            {"email": email},
            {"$inc": {"generation_count": 1}}
        )
        
        # Get the updated user data
        user = collection.find_one({"email": email})
        return user.get("generation_count", 0)
    except Exception as e:
        st.error(f"Erro ao atualizar contagem de gerações: {e}")
        return None

# Function to check if user has reached generation limit
def check_generation_limit(email, collection, limit=2):
    user = collection.find_one({"email": email})
    if user:
        generation_count = user.get("generation_count", 0)
        return generation_count >= limit
    return False

    # Main function
def main():
    st.set_page_config(page_title="Gerador de Currículo", page_icon="📄", layout="wide")
    st.title("📄 Gerador de Currículo Personalizado")

    # Initialize session state for user authentication and generation count
    if "signed_in" not in st.session_state:
        st.session_state.signed_in = False
    if "generation_count" not in st.session_state:
        st.session_state.generation_count = 0

    # Connect to MongoDB
    collection = connect_to_mongodb()
    if collection is None:
        return

    # Sign-up/Sign-in form
    if not st.session_state.signed_in:
        st.header("Cadastro/Login")
        email = st.text_input("Email")
        password = st.text_input("Senha", type="password")
        if st.button("Entrar/Cadastrar"):
            if email and password:
                if check_user_exists(email, collection):
                    user = authenticate_user(email, password, collection)
                    if user:
                        st.session_state.signed_in = True
                        st.session_state.email = email
                        st.session_state.generation_count = user.get("generation_count", 0)
                        st.success("Login bem-sucedido!")
                        st.rerun()
                    else:
                        st.error("Email ou senha incorretos.")
                else:
                    add_user(email, password, collection)
                    st.session_state.signed_in = True
                    st.session_state.email = email
                    st.session_state.generation_count = 0
                    st.success("Cadastro bem-sucedido!")
                    st.rerun()
            else:
                st.error("Por favor, insira email e senha.")
        return

    # Main content for signed-in users
    st.write(f"Bem-vindo, {st.session_state.email}!")
    
    # Check generation limit
    has_reached_limit = check_generation_limit(st.session_state.email, collection)
    if has_reached_limit:
        st.warning("⚠️ Você atingiu o limite de 2 currículos gerados. Para gerar mais currículos, por favor, realize um pagamento. Pague uma taxa de 200 MTS para o número 876513064 (Ernestina Jose).")
        # st.button("Realizar Pagamento", type="primary")
        st.info("Entre em contato conosco para mais informações sobre pagamentos.")
        return
    
    # Display remaining generations
    st.info(f"Você tem {2 - st.session_state.generation_count} gerações de currículo restantes em sua conta gratuita.")
    st.write("Preencha os detalhes abaixo para gerar seu currículo.")

    # List available templates
    templates, templates_dir = list_templates()
    if not templates:
        st.warning("Nenhum template encontrado na pasta 'templates'. Adicione templates .docx para continuar.")
        return

    # Let the user select a template
    selected_template = st.selectbox("Escolha um template", templates)

    # Sidebar for user inputs
    with st.sidebar:
        st.header("Informações Pessoais")
        name = st.text_input("Nome Completo")
        phone = st.text_input("Telefone")
        linkedin = st.text_input("LinkedIn (opcional)")
        languages = st.text_input("Idiomas (separados por vírgula)", "Português, Inglês")
        interests = st.text_input("Atividades e Interesses (separados por vírgula)", "Teatro, Conservação Ambiental, Artes")

        st.header("Detalhes Profissionais")
        industry = st.selectbox("Indústria", ["Tecnologia", "Saúde", "Educação", "Finanças", "Marketing"])
        job_type = st.selectbox("Tipo de Emprego", ["Desenvolvedor", "Analista", "Gerente", "Consultor"])

    # Main content area
    with st.container():
        st.header("Experiência Profissional")
        experiences = []
        num_experiences = st.number_input("Quantas experiências profissionais você tem?", min_value=1, value=1)
        for i in range(num_experiences):
            st.write(f"Experiência {i+1}")
            company = st.text_input(f"Nome da Empresa {i+1}")
            title = st.text_input(f"Cargo {i+1}")
            start_date = st.text_input(f"Data de Início {i+1} (MM/AAAA)")
            end_date = st.text_input(f"Data de Término {i+1} (MM/AAAA)")
            duties = st.text_area(f"Responsabilidades {i+1}")
            experiences.append(f"{title} na {company} ({start_date} - {end_date}): {duties}")

        st.header("Educação")
        educations = []
        num_educations = st.number_input("Quantas formações acadêmicas você tem?", min_value=1, value=1)
        for i in range(num_educations):
            st.write(f"Formação {i+1}")
            degree = st.text_input(f"Grau {i+1} (ex: Bacharel em Administração)")
            institution = st.text_input(f"Instituição {i+1}")
            graduation_date = st.text_input(f"Data de Conclusão {i+1} (MM/AAAA)")
            educations.append(f"{degree} na {institution} ({graduation_date})")

        st.header("Habilidades")
        skills = st.text_area("Liste suas habilidades (separadas por vírgula)", "Contabilidade, PDV, Comunicação")

    # Generate Resume
    if st.button("Gerar Currículo"):
        if not client.api_key:
            st.error("Por favor, configure sua chave da API da OpenAI.")
        else:
            with st.spinner("Gerando seu currículo..."):
                # Format experiences and educations
                experiences_formatted = "\n".join([f"- {exp}" for exp in experiences])
                educations_formatted = "\n".join([f"- {edu}" for edu in educations])

                # Generate resume content
                resume_content = generate_resume(
                    name, st.session_state.email, phone, industry, job_type,
                    experiences_formatted, educations_formatted, skills.split(","), languages.split(","), linkedin
                )

                # Update generation count in the database
                new_count = update_generation_count(st.session_state.email, collection)
                if new_count is not None:
                    st.session_state.generation_count = new_count

                # Display Resume
                st.subheader("Seu Currículo Gerado")
                st.text(resume_content)

                # Create Word document
                doc_bytes = create_word_doc(resume_content, name, st.session_state.email, phone, linkedin, selected_template, templates_dir)
                if doc_bytes:
                    st.download_button(
                        label="Baixar Currículo em Word",
                        data=doc_bytes,
                        file_name="curriculo.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )

                # Show warning if user reached limit after this generation
                if st.session_state.generation_count >= 2:
                    st.warning("⚠️ Você atingiu o limite de currículos gratuitos. Para gerar mais currículos, por favor, realize um pagamento. Pague uma taxa de 200 MTS para o número 876513064 (Ernestina Jose).")
                    # st.button("Realizar Pagamento", type="primary")
                else:
                    st.info(f"Você tem {2 - st.session_state.generation_count} gerações de currículo restantes em sua conta gratuita.")

if __name__ == "__main__":
    main()